"""Convert a markdown file to a .docx document.

Handles: headings, paragraphs, bold/italic/code inline, code blocks,
tables (pipe syntax), bullet/numbered lists, horizontal rules, links.

Usage:
    python scripts/md_to_docx.py input.md output.docx
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


INLINE_RE = re.compile(
    r"(\*\*([^*]+)\*\*|\*([^*]+)\*|`([^`]+)`|\[([^\]]+)\]\(([^)]+)\))"
)


def add_inline_runs(paragraph, text):
    """Parse a string for **bold**, *italic*, `code`, [link](url) and add as runs."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        full, bold, italic, code, link_text, _link_url = m.groups()
        if bold is not None:
            r = paragraph.add_run(bold)
            r.bold = True
        elif italic is not None:
            r = paragraph.add_run(italic)
            r.italic = True
        elif code is not None:
            r = paragraph.add_run(code)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x8B, 0x2E, 0x2E)
        elif link_text is not None:
            r = paragraph.add_run(link_text)
            r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            r.underline = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def parse_table_block(lines):
    """Lines are like '| a | b |'. Returns list of rows, each a list of cells."""
    rows = []
    for raw in lines:
        s = raw.strip()
        if not s.startswith("|"):
            continue
        if re.match(r"^\|[\s\-:|]+\|$", s):
            continue
        cells = [c.strip() for c in s.strip("|").split("|")]
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, val)
            if i == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "1F4E79")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    text = "\n".join(code_lines)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def md_to_docx(md_path: Path, docx_path: Path):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    raw = md_path.read_text(encoding="utf-8")
    lines = raw.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_block(table_lines)
            add_table(doc, rows)
            continue

        if stripped in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            h = doc.add_heading(level=min(level, 4))
            add_inline_runs(h, text)
            i += 1
            continue

        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:])
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_inline_runs(p, " ".join(quote_lines))
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|```|\||---|>\s|\s*[-*+]\s|\s*\d+\.\s)", lines[i]
        ):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(para_lines))

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_to_docx(Path(sys.argv[1]), Path(sys.argv[2]))
